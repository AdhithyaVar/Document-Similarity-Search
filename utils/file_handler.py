"""
File Handler Module
Handles extraction of text from various document formats
"""

import fitz  # PyMuPDF
import docx
import pandas as pd
from typing import Optional, BinaryIO
import logging
from pathlib import Path
import re

from config.settings import (
    MAX_FILE_SIZE,
    SUPPORTED_FORMATS,
    VALIDATION_RULES,
    MESSAGES,
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class FileHandler:
    """
    Handles file operations and text extraction from multiple formats
    """

    def __init__(self):
        """Initialize FileHandler"""
        self.supported_formats = SUPPORTED_FORMATS
        self.max_file_size = MAX_FILE_SIZE

    def validate_file(self, file: BinaryIO) -> tuple[bool, Optional[str]]:
        """
        Validate uploaded file
        
        Args:
            file: Uploaded file object
            
        Returns:
            tuple: (is_valid, error_message)
        """
        try:
            # Check file size
            file.seek(0, 2)  # Seek to end
            size = file.tell()
            file.seek(0)  # Reset to beginning
            
            if size < VALIDATION_RULES["min_file_size"]:
                return False, MESSAGES["file_too_small"]
            
            if size > VALIDATION_RULES["max_file_size"]:
                return False, MESSAGES["file_too_large"]
            
            # Check file extension
            file_ext = Path(file.name).suffix.lower().lstrip(".")
            if file_ext not in VALIDATION_RULES["allowed_extensions"]:
                return False, MESSAGES["invalid_format"]
            
            # Check for forbidden patterns
            for pattern in VALIDATION_RULES["forbidden_patterns"]:
                if re.search(pattern, file.name, re.IGNORECASE):
                    return False, "Forbidden file type detected"
            
            return True, None
            
        except Exception as e:
            logger.error(f"File validation error: {str(e)}")
            return False, f"Validation error: {str(e)}"

    def extract_text(self, file: BinaryIO) -> Optional[str]:
        """
        Extract text from uploaded file
        
        Args:
            file: Uploaded file object
            
        Returns:
            Extracted text or None if extraction fails
        """
        try:
            # Validate file first
            is_valid, error_msg = self.validate_file(file)
            if not is_valid:
                logger.error(f"File validation failed: {error_msg}")
                return None
            
            # Get file extension
            file_ext = Path(file.name).suffix.lower().lstrip(".")
            
            # Route to appropriate extractor
            extractors = {
                "pdf": self._extract_from_pdf,
                "docx": self._extract_from_docx,
                "txt": self._extract_from_txt,
                "csv": self._extract_from_csv,
                "xlsx": self._extract_from_excel,
            }
            
            extractor = extractors.get(file_ext)
            if not extractor:
                logger.error(f"No extractor for format: {file_ext}")
                return None
            
            # Reset file pointer
            file.seek(0)
            
            # Extract text
            text = extractor(file)
            
            if text and len(text.strip()) > 0:
                logger.info(f"Successfully extracted {len(text)} characters from {file.name}")
                return text
            else:
                logger.warning(f"No text extracted from {file.name}")
                return None
                
        except Exception as e:
            logger.error(f"Text extraction error: {str(e)}")
            return None

    def _extract_from_pdf(self, file: BinaryIO) -> str:
        """
        Extract text from PDF file
        
        Args:
            file: PDF file object
            
        Returns:
            Extracted text
        """
        try:
            # Open PDF
            pdf_bytes = file.read()
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            
            text_parts = []
            
            # Extract text from each page
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Try different extraction methods
                text = page.get_text("text")
                
                # If text extraction fails, try blocks
                if not text or len(text.strip()) < 10:
                    blocks = page.get_text("blocks")
                    text = "\n".join([block[4] for block in blocks if len(block) > 4])
                
                text_parts.append(text)
            
            doc.close()
            
            full_text = "\n\n".join(text_parts)
            return self._clean_text(full_text)
            
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise

    def _extract_from_docx(self, file: BinaryIO) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file: DOCX file object
            
        Returns:
            Extracted text
        """
        try:
            doc = docx.Document(file)
            
            text_parts = []
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            full_text = "\n\n".join(text_parts)
            return self._clean_text(full_text)
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            raise

    def _extract_from_txt(self, file: BinaryIO) -> str:
        """
        Extract text from TXT file
        
        Args:
            file: TXT file object
            
        Returns:
            Extracted text
        """
        try:
            # Try multiple encodings
            encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
            
            for encoding in encodings:
                try:
                    file.seek(0)
                    text = file.read().decode(encoding)
                    return self._clean_text(text)
                except UnicodeDecodeError:
                    continue
            
            # If all fail, use UTF-8 with error handling
            file.seek(0)
            text = file.read().decode("utf-8", errors="ignore")
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"TXT extraction error: {str(e)}")
            raise

    def _extract_from_csv(self, file: BinaryIO) -> str:
        """
        Extract text from CSV file
        
        Args:
            file: CSV file object
            
        Returns:
            Extracted text
        """
        try:
            df = pd.read_csv(file, encoding="utf-8", on_bad_lines="skip")
            
            # Convert DataFrame to readable text
            text_parts = []
            
            # Add headers
            text_parts.append(" | ".join(df.columns))
            text_parts.append("-" * 50)
            
            # Add rows
            for _, row in df.iterrows():
                row_text = " | ".join([str(val) for val in row.values])
                text_parts.append(row_text)
            
            full_text = "\n".join(text_parts)
            return self._clean_text(full_text)
            
        except Exception as e:
            logger.error(f"CSV extraction error: {str(e)}")
            raise

    def _extract_from_excel(self, file: BinaryIO) -> str:
        """
        Extract text from Excel file
        
        Args:
            file: Excel file object
            
        Returns:
            Extracted text
        """
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file)
            text_parts = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                # Add sheet name
                text_parts.append(f"\n=== Sheet: {sheet_name} ===\n")
                
                # Add headers
                text_parts.append(" | ".join(df.columns))
                text_parts.append("-" * 50)
                
                # Add rows
                for _, row in df.iterrows():
                    row_text = " | ".join([str(val) for val in row.values])
                    text_parts.append(row_text)
            
            full_text = "\n".join(text_parts)
            return self._clean_text(full_text)
            
        except Exception as e:
            logger.error(f"Excel extraction error: {str(e)}")
            raise

    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text
        
        Args:
            text: Raw text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Normalize line breaks
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text

    def get_file_info(self, file: BinaryIO) -> dict:
        """
        Get metadata about uploaded file
        
        Args:
            file: File object
            
        Returns:
            Dictionary with file information
        """
        try:
            file.seek(0, 2)
            size = file.tell()
            file.seek(0)
            
            return {
                "name": file.name,
                "size": size,
                "size_mb": round(size / (1024 * 1024), 2),
                "extension": Path(file.name).suffix.lower().lstrip("."),
                "is_valid": self.validate_file(file)[0],
            }
            
        except Exception as e:
            logger.error(f"Error getting file info: {str(e)}")
            return {}